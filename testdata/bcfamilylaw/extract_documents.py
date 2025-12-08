"""
Extract documents from benchmark JSON into various file formats for testing.

This script reads the benchmark JSON and creates:
- PDF files (using reportlab)
- DOCX files (using python-docx)
- TXT files (plain text)

Then updates the JSON to reference file paths instead of inline content.
"""

import json
import os
from pathlib import Path

# Check for required libraries
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.units import inch
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False
    print("Warning: reportlab not installed. PDF generation will be skipped.")
    print("Install with: pip install reportlab")

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Warning: python-docx not installed. DOCX generation will be skipped.")
    print("Install with: pip install python-docx")


def create_txt_file(content: str, filepath: Path) -> bool:
    """Create a plain text file."""
    try:
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        return True
    except Exception as e:
        print(f"Error creating TXT file {filepath}: {e}")
        return False


def create_pdf_file(content: str, filepath: Path, title: str) -> bool:
    """Create a PDF file using reportlab."""
    if not HAS_REPORTLAB:
        return False

    try:
        doc = SimpleDocTemplate(
            str(filepath),
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )

        styles = getSampleStyleSheet()

        # Create custom styles for legal documents
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=14,
            spaceAfter=20,
            alignment=1  # Center
        )

        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontSize=10,
            leading=14,
            spaceAfter=6
        )

        story = []

        # Add title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.25 * inch))

        # Split content into paragraphs and add to story
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                # Replace single newlines with <br/> for proper line breaks
                para_html = para.replace('\n', '<br/>')
                # Escape special characters
                para_html = para_html.replace('&', '&amp;')
                story.append(Paragraph(para_html, body_style))
                story.append(Spacer(1, 0.1 * inch))

        doc.build(story)
        return True
    except Exception as e:
        print(f"Error creating PDF file {filepath}: {e}")
        return False


def create_docx_file(content: str, filepath: Path, title: str) -> bool:
    """Create a Word document using python-docx."""
    if not HAS_DOCX:
        return False

    try:
        doc = Document()

        # Add title
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add content paragraphs
        paragraphs = content.split('\n\n')
        for para in paragraphs:
            if para.strip():
                p = doc.add_paragraph()
                # Handle line breaks within paragraphs
                lines = para.split('\n')
                for i, line in enumerate(lines):
                    if i > 0:
                        p.add_run('\n')
                    p.add_run(line)

        doc.save(str(filepath))
        return True
    except Exception as e:
        print(f"Error creating DOCX file {filepath}: {e}")
        return False


def main():
    # Get the directory of this script
    script_dir = Path(__file__).parent
    benchmark_file = script_dir / "bc_family_law_benchmark_with_files.json"
    documents_dir = script_dir / "documents"

    # Create documents directory
    documents_dir.mkdir(exist_ok=True)

    # Load benchmark JSON
    with open(benchmark_file, 'r', encoding='utf-8') as f:
        benchmark = json.load(f)

    # Define format assignments for each document
    # Mix of formats to test different upload capabilities
    format_assignments = {
        "DOC-001": "pdf",   # Court Order -> PDF (common for court documents)
        "DOC-002": "docx",  # Financial Statement -> Word (often filled forms)
        "DOC-003": "pdf",   # Separation Agreement -> PDF (signed documents)
        "DOC-004": "docx",  # Affidavit -> Word (drafts often in Word)
        "DOC-005": "txt",   # Provincial Court Order -> TXT (for variety)
        "DOC-006": "docx",  # Financial Statement -> Word (Chen matter second doc)
    }

    print(f"Extracting {len(benchmark['documents'])} documents...")
    print(f"Output directory: {documents_dir}")
    print()

    # Process each document
    for doc in benchmark['documents']:
        doc_id = doc['doc_id']
        title = doc['title']
        content = doc['content']
        doc_type = doc['doc_type']

        # Get assigned format
        file_format = format_assignments.get(doc_id, 'txt')

        # Create filename (sanitize title)
        safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()
        safe_title = safe_title.replace(' ', '_')
        filename = f"{doc_id}_{safe_title}.{file_format}"
        filepath = documents_dir / filename

        print(f"Creating {doc_id}: {filename}")
        print(f"  Type: {doc_type}")
        print(f"  Format: {file_format.upper()}")

        # Create the file based on format
        success = False
        if file_format == 'pdf':
            success = create_pdf_file(content, filepath, title)
            if not success and HAS_DOCX:
                # Fallback to DOCX if PDF fails
                filename = filename.replace('.pdf', '.docx')
                filepath = documents_dir / filename
                success = create_docx_file(content, filepath, title)
                file_format = 'docx'
        elif file_format == 'docx':
            success = create_docx_file(content, filepath, title)
            if not success:
                # Fallback to TXT
                filename = filename.replace('.docx', '.txt')
                filepath = documents_dir / filename
                success = create_txt_file(content, filepath)
                file_format = 'txt'
        else:  # txt
            success = create_txt_file(content, filepath)

        if success:
            print(f"  Created: {filepath.name}")
            # Update document in benchmark with file path
            doc['file_path'] = str(filepath.relative_to(script_dir))
            doc['file_format'] = file_format
        else:
            print(f"  FAILED to create file")
            # Keep content in JSON as fallback
            doc['file_path'] = None
            doc['file_format'] = None

        print()

    # Save updated benchmark JSON
    with open(benchmark_file, 'w', encoding='utf-8') as f:
        json.dump(benchmark, f, indent=2)

    print(f"Updated benchmark saved to: {benchmark_file}")
    print()

    # Print summary
    print("=" * 60)
    print("SUMMARY")
    print("=" * 60)
    formats_created = {}
    for doc in benchmark['documents']:
        fmt = doc.get('file_format', 'none')
        formats_created[fmt] = formats_created.get(fmt, 0) + 1

    for fmt, count in formats_created.items():
        print(f"  {fmt.upper()}: {count} file(s)")

    print()
    print("Files are ready in:", documents_dir)
    print()
    print("Next step: Run the upload script to test the API")
    print("  python upload_documents.py")


if __name__ == "__main__":
    main()
