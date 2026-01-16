"""
Text extraction service for legal documents.

Extracts text from PDF, DOCX, TXT, and image files for RAG pipeline processing.
Optimized for legal documents with proper handling of tables, headers, and formatting.

Includes OCR support via AWS Textract/Tesseract for scanned PDFs and images.
"""

import io
import logging
from typing import Dict, Optional, Tuple
from pathlib import Path

import pdfplumber
from docx import Document
import charset_normalizer

from app.services.ocr import ocr_service, is_likely_scanned_pdf

logger = logging.getLogger(__name__)


class TextExtractionService:
    """Service for extracting text from various document formats."""

    # Maximum file size to process (500 MB for large SR&ED documents)
    MAX_FILE_SIZE = 500 * 1024 * 1024

    # Supported image extensions for OCR
    IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.tiff', '.tif'}

    @staticmethod
    def extract_text(
        file_content: bytes,
        filename: str,
        mime_type: Optional[str] = None,
        storage_path: Optional[str] = None
    ) -> Dict:
        """
        Extract text from document based on file type.

        For PDFs, automatically detects scanned documents and uses OCR if needed.

        Args:
            file_content: Raw file bytes
            filename: Original filename (used to determine file type)
            mime_type: MIME type of the file (optional)
            storage_path: S3 storage path (required for OCR on scanned PDFs)

        Returns:
            Dictionary with extraction results:
            {
                'success': bool,
                'text': str,
                'metadata': {
                    'page_count': int,
                    'word_count': int,
                    'char_count': int,
                    'extraction_method': str,
                    'ocr_applied': bool,
                    'ocr_engine': str (if OCR applied),
                    'ocr_pages_processed': int (if OCR applied),
                    'ocr_confidence_avg': float (if OCR applied)
                },
                'error': str (if success=False)
            }
        """
        # Check file size
        if len(file_content) > TextExtractionService.MAX_FILE_SIZE:
            return {
                'success': False,
                'text': '',
                'metadata': {},
                'error': f'File too large: {len(file_content)} bytes (max: {TextExtractionService.MAX_FILE_SIZE})'
            }

        # Determine file type
        file_ext = Path(filename).suffix.lower()
        logger.info(f"Text extraction for file '{filename}' with extension '{file_ext}'")

        try:
            if file_ext == '.pdf':
                return TextExtractionService._extract_pdf(file_content, storage_path)
            elif file_ext in ['.docx', '.doc']:
                return TextExtractionService._extract_docx(file_content)
            elif file_ext == '.txt':
                return TextExtractionService._extract_txt(file_content)
            elif file_ext in TextExtractionService.IMAGE_EXTENSIONS:
                logger.info(f"Routing to image extraction for {filename}")
                return TextExtractionService._extract_image(file_content, storage_path)
            else:
                logger.warning(f"Unsupported file type: {file_ext} (supported images: {TextExtractionService.IMAGE_EXTENSIONS})")
                return {
                    'success': False,
                    'text': '',
                    'metadata': {},
                    'error': f'Unsupported file type: {file_ext}'
                }

        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {str(e)}", exc_info=True)
            return {
                'success': False,
                'text': '',
                'metadata': {},
                'error': f'Extraction error: {str(e)}'
            }

    @staticmethod
    def _extract_pdf(file_content: bytes, storage_path: Optional[str] = None) -> Dict:
        """
        Extract text from PDF using pdfplumber.

        If the PDF appears to be scanned (insufficient extractable text),
        automatically falls back to OCR using AWS Textract.

        pdfplumber is better than PyPDF2 for:
        - Table extraction
        - Layout preservation
        - Text quality in complex documents
        """
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                pages_text = []
                page_count = len(pdf.pages)

                for page_num, page in enumerate(pdf.pages, start=1):
                    try:
                        # Extract text with layout preservation
                        text = page.extract_text()
                        if text:
                            # Add page marker for citation tracking
                            pages_text.append(f"[Page {page_num}]\n{text}")
                    except Exception as e:
                        logger.warning(f"Failed to extract page {page_num}: {str(e)}")
                        pages_text.append(f"[Page {page_num}]\n[Extraction failed]")

                full_text = "\n\n".join(pages_text)

                # Check if this appears to be a scanned PDF
                if is_likely_scanned_pdf(full_text, page_count):
                    logger.info(f"PDF appears to be scanned, attempting OCR extraction")

                    # Use unified OCR service (tries Textract, falls back to Tesseract)
                    return TextExtractionService._extract_pdf_with_ocr(
                        file_content, storage_path, page_count
                    )

                # Standard extraction succeeded
                word_count = len(full_text.split())
                char_count = len(full_text)

                return {
                    'success': True,
                    'text': full_text,
                    'metadata': {
                        'page_count': page_count,
                        'word_count': word_count,
                        'char_count': char_count,
                        'extraction_method': 'pdfplumber',
                        'ocr_applied': False
                    },
                    'error': None
                }

        except Exception as e:
            logger.error(f"PDF extraction error: {str(e)}")
            return {
                'success': False,
                'text': '',
                'metadata': {},
                'error': f'PDF extraction failed: {str(e)}'
            }

    @staticmethod
    def _extract_pdf_with_ocr(
        file_content: bytes,
        storage_path: Optional[str],
        page_count: int
    ) -> Dict:
        """
        Extract text from a scanned PDF using OCR.

        Uses the unified OCR service which tries Textract first (production),
        then falls back to Tesseract (local development).

        Args:
            file_content: Raw PDF bytes (needed for Tesseract fallback)
            storage_path: S3 key of the document (needed for Textract)
            page_count: Number of pages (from pdfplumber)

        Returns:
            Dictionary with extraction results including OCR metadata
        """
        import asyncio

        try:
            # Run async OCR extraction
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        try:
            ocr_result = loop.run_until_complete(
                ocr_service.extract_text(file_content, storage_path)
            )

            if ocr_result['success']:
                full_text = ocr_result['text']
                word_count = len(full_text.split())
                char_count = len(full_text)
                ocr_engine = ocr_result.get('ocr_engine', 'unknown')

                return {
                    'success': True,
                    'text': full_text,
                    'metadata': {
                        'page_count': ocr_result['pages_processed'] or page_count,
                        'word_count': word_count,
                        'char_count': char_count,
                        'extraction_method': f'pdfplumber_with_{ocr_engine}',
                        'ocr_applied': True,
                        'ocr_engine': ocr_engine,
                        'ocr_pages_processed': ocr_result['pages_processed'],
                        'ocr_confidence_avg': ocr_result['confidence_avg']
                    },
                    'error': None
                }
            else:
                # OCR failed, return error
                logger.error(f"OCR extraction failed: {ocr_result['error']}")
                return {
                    'success': False,
                    'text': '',
                    'metadata': {
                        'page_count': page_count,
                        'ocr_applied': True,
                        'ocr_engine': ocr_result.get('ocr_engine')
                    },
                    'error': f"OCR extraction failed: {ocr_result['error']}"
                }

        except Exception as e:
            logger.error(f"OCR extraction error: {str(e)}", exc_info=True)
            return {
                'success': False,
                'text': '',
                'metadata': {'page_count': page_count},
                'error': f'OCR extraction failed: {str(e)}'
            }

    @staticmethod
    def _extract_image(file_content: bytes, storage_path: Optional[str] = None) -> Dict:
        """
        Extract text from image files using OCR.

        Images always require OCR - no text extraction fallback.

        Args:
            file_content: Raw image bytes (PNG, JPG, JPEG, TIFF)
            storage_path: S3 key of the document (needed for Textract)

        Returns:
            Dictionary with extraction results including OCR metadata
        """
        from app.services.ocr import tesseract_ocr_service, textract_ocr_service
        import asyncio

        logger.info(f"Extracting text from image using OCR (file size: {len(file_content)} bytes)")

        ocr_result = None
        ocr_engine = None

        # Check OCR availability
        textract_available = textract_ocr_service.is_available()
        tesseract_available = tesseract_ocr_service.is_available()
        logger.info(f"OCR availability - Textract: {textract_available}, Tesseract: {tesseract_available}")

        # Try Textract first if available (production)
        if storage_path and textract_available:
            logger.info("Attempting image OCR with AWS Textract")
            try:
                # Handle async Textract call
                try:
                    loop = asyncio.get_event_loop()
                    if loop.is_running():
                        # We're in an async context, create a new thread to run it
                        import concurrent.futures
                        with concurrent.futures.ThreadPoolExecutor() as executor:
                            future = executor.submit(
                                asyncio.run,
                                textract_ocr_service.extract_text_from_s3_document(storage_path)
                            )
                            ocr_result = future.result()
                    else:
                        ocr_result = loop.run_until_complete(
                            textract_ocr_service.extract_text_from_s3_document(storage_path)
                        )
                except RuntimeError:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    ocr_result = loop.run_until_complete(
                        textract_ocr_service.extract_text_from_s3_document(storage_path)
                    )

                if ocr_result and ocr_result.get('success'):
                    ocr_engine = 'textract'
                else:
                    logger.warning(f"Textract failed for image, falling back to Tesseract: {ocr_result.get('error') if ocr_result else 'Unknown error'}")
                    ocr_result = None
            except Exception as e:
                logger.warning(f"Textract failed for image: {e}, falling back to Tesseract")
                ocr_result = None

        # Fall back to Tesseract (local development or Textract failure)
        if ocr_result is None and tesseract_available:
            logger.info("Using Tesseract OCR for image text extraction")
            try:
                ocr_result = tesseract_ocr_service.extract_text_from_image_bytes(file_content)
                logger.info(f"Tesseract OCR result: success={ocr_result.get('success')}, error={ocr_result.get('error')}")
                if ocr_result.get('success'):
                    ocr_engine = 'tesseract'
            except Exception as e:
                logger.error(f"Tesseract OCR threw exception: {e}", exc_info=True)
                ocr_result = {
                    'success': False,
                    'text': '',
                    'pages_processed': 0,
                    'confidence_avg': None,
                    'error': f'Tesseract exception: {str(e)}'
                }

        # Handle no OCR available
        if ocr_result is None:
            error_msg = 'No OCR engine available. Install Tesseract or configure AWS credentials.'
            logger.error(error_msg)
            return {
                'success': False,
                'text': '',
                'metadata': {'page_count': 1, 'ocr_applied': True},
                'error': error_msg
            }

        # Process OCR result
        if ocr_result['success']:
            full_text = ocr_result['text']
            word_count = len(full_text.split())
            char_count = len(full_text)

            return {
                'success': True,
                'text': full_text,
                'metadata': {
                    'page_count': 1,  # Images are single-page
                    'word_count': word_count,
                    'char_count': char_count,
                    'extraction_method': f'ocr_{ocr_engine}',
                    'ocr_applied': True,
                    'ocr_engine': ocr_engine,
                    'ocr_pages_processed': 1,
                    'ocr_confidence_avg': ocr_result['confidence_avg']
                },
                'error': None
            }
        else:
            # OCR failed
            logger.error(f"Image OCR extraction failed: {ocr_result['error']}")
            return {
                'success': False,
                'text': '',
                'metadata': {
                    'page_count': 1,
                    'ocr_applied': True,
                    'ocr_engine': ocr_engine
                },
                'error': f"Image OCR extraction failed: {ocr_result['error']}"
            }

    @staticmethod
    def _extract_docx(file_content: bytes) -> Dict:
        """
        Extract text from DOCX files using python-docx.

        Extracts:
        - Paragraph text
        - Table contents
        - Headers and footers
        """
        try:
            doc = Document(io.BytesIO(file_content))
            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells]
                    table_text.append(" | ".join(row_text))

                if table_text:
                    text_parts.append("\n[Table]\n" + "\n".join(table_text) + "\n[End Table]")

            full_text = "\n\n".join(text_parts)
            word_count = len(full_text.split())
            char_count = len(full_text)

            # Estimate page count (rough approximation: 500 words per page)
            page_count = max(1, word_count // 500)

            return {
                'success': True,
                'text': full_text,
                'metadata': {
                    'page_count': page_count,
                    'word_count': word_count,
                    'char_count': char_count,
                    'extraction_method': 'python-docx'
                },
                'error': None
            }

        except Exception as e:
            logger.error(f"DOCX extraction error: {str(e)}")
            return {
                'success': False,
                'text': '',
                'metadata': {},
                'error': f'DOCX extraction failed: {str(e)}'
            }

    @staticmethod
    def _extract_txt(file_content: bytes) -> Dict:
        """
        Extract text from plain text files with encoding detection.

        Uses charset_normalizer for robust encoding detection.
        """
        try:
            # Detect encoding
            detection_result = charset_normalizer.from_bytes(file_content).best()

            if detection_result is None:
                # Fallback to UTF-8
                text = file_content.decode('utf-8', errors='ignore')
            else:
                text = str(detection_result)

            word_count = len(text.split())
            char_count = len(text)

            # Estimate page count (rough approximation: 500 words per page)
            page_count = max(1, word_count // 500)

            return {
                'success': True,
                'text': text,
                'metadata': {
                    'page_count': page_count,
                    'word_count': word_count,
                    'char_count': char_count,
                    'extraction_method': 'charset_normalizer'
                },
                'error': None
            }

        except Exception as e:
            logger.error(f"TXT extraction error: {str(e)}")
            return {
                'success': False,
                'text': '',
                'metadata': {},
                'error': f'TXT extraction failed: {str(e)}'
            }

    @staticmethod
    def extract_metadata_only(
        file_content: bytes,
        filename: str
    ) -> Dict:
        """
        Extract only metadata without full text extraction.
        Useful for quick file analysis.

        Returns:
            Dictionary with metadata:
            {
                'page_count': int,
                'file_size': int,
                'file_type': str
            }
        """
        file_ext = Path(filename).suffix.lower()
        metadata = {
            'file_size': len(file_content),
            'file_type': file_ext
        }

        try:
            if file_ext == '.pdf':
                with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                    metadata['page_count'] = len(pdf.pages)
            elif file_ext in ['.docx', '.doc']:
                doc = Document(io.BytesIO(file_content))
                # Rough page count estimate
                para_count = len(doc.paragraphs)
                metadata['page_count'] = max(1, para_count // 30)  # ~30 paragraphs per page
            elif file_ext == '.txt':
                text = file_content.decode('utf-8', errors='ignore')
                word_count = len(text.split())
                metadata['page_count'] = max(1, word_count // 500)
            elif file_ext in TextExtractionService.IMAGE_EXTENSIONS:
                # Images are always single-page
                metadata['page_count'] = 1
            else:
                metadata['page_count'] = 0

            return metadata

        except Exception as e:
            logger.error(f"Metadata extraction failed for {filename}: {str(e)}")
            return metadata


# Create singleton instance
text_extraction_service = TextExtractionService()
